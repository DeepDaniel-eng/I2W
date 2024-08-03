import streamlit as st
import os
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import tempfile

def create_word_document(files, images_per_row=3, image_width_inch=2, image_height_inch=2, padding_inch=0.1):
    doc = Document()

    table = doc.add_table(rows=0, cols=images_per_row)
    table.autofit = False

    row_cells = None
    for i, file in enumerate(files):
        if i % images_per_row == 0:
            row_cells = table.add_row().cells

        paragraph = row_cells[i % images_per_row].paragraphs[0]
        run = paragraph.add_run()
        
        # Load image data from the uploaded file
        image_stream = BytesIO(file.getvalue())
        run.add_picture(image_stream, width=Inches(image_width_inch), height=Inches(image_height_inch))

        # Set cell width
        cell_width = Inches(image_width_inch + 2 * padding_inch)
        row_cells[i % images_per_row].width = cell_width

        # Apply padding
        for cell in row_cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(padding_inch * 72)
            cell.paragraphs[0].paragraph_format.space_after = Pt(padding_inch * 72)

    output_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(output_file.name)
    output_file.close()
    return output_file.name

st.title("Image to Word Document Generator")

uploaded_files = st.file_uploader("Upload Images", type=['png', 'jpg', 'jpeg', 'gif'], accept_multiple_files=True)

if st.button("Generate Word Document"):
    if uploaded_files:
        output_path = create_word_document(uploaded_files)
        with open(output_path, "rb") as f:
            st.download_button("Download Word Document", f, file_name="output.docx")
        st.success("Word document created successfully!")
    else:
        st.error("Please upload at least one image.")
